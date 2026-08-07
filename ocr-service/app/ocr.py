"""OCR & text extraction module for mediation application forms.

Pipeline:
  1. Extract text from uploaded file:
     - PDF: pdfplumber (electronic) → fallback pdf2image + PaddleOCR (scanned)
     - Word (.docx): python-docx
     - Images: PaddleOCR
     - .txt/.md: direct read
   2. Local LLM structured extraction (OpenAI-compatible, default local server)
      → JSON fields for form autofill.

Graceful degradation: if PaddleOCR/pdfplumber not installed, returns error hint.
"""
from __future__ import annotations

import io
import os
import time
from pathlib import Path
from typing import Any, Optional

# ── Lazy-loaded engines ─────────────────────────────────────
# Uses RapidOCR (PaddleOCR models + onnxruntime inference, no paddlepaddle dependency).
# paddlepaddle has no prebuilt wheels for macOS arm64, so RapidOCR is the cross-platform fallback.
_ocr_engine = None
_ocr_error: Optional[str] = None


def _get_ocr_engine():
    """Lazy-load RapidOCR engine (once)."""
    global _ocr_engine, _ocr_error
    if _ocr_engine is not None or _ocr_error is not None:
        return _ocr_engine

    try:
        from rapidocr_onnxruntime import RapidOCR
        t0 = time.time()
        _ocr_engine = RapidOCR()
        print(f"[OCR] ✅ RapidOCR loaded in {time.time() - t0:.1f}s")
    except Exception as e:
        _ocr_error = str(e)
        print(f"[OCR] ⚠️  RapidOCR unavailable: {e}")
        print(f"[OCR]    Install with: pip install rapidocr-onnxruntime")
    return _ocr_engine


# ── Text extraction by file type ────────────────────────────
def _extract_pdf_text(file_bytes: bytes) -> tuple[str, bool]:
    """Extract text from PDF. Returns (text, used_ocr).
    Tries pdfplumber first; if text is too short, falls back to PaddleOCR on rendered images.
    """
    # Try electronic text extraction
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n"
    except Exception as e:
        print(f"[OCR] pdfplumber failed: {e}")

    # If we got enough text, treat as electronic PDF
    if len(text.strip()) >= 50:
        return text.strip(), False

    # Fallback: render pages to images and OCR
    print("[OCR] PDF text too short, falling back to image OCR...")
    ocr_engine = _get_ocr_engine()
    if ocr_engine is None:
        return text.strip(), False

    try:
        from pdf2image import convert_from_bytes
        images = convert_from_bytes(file_bytes, dpi=200)
        ocr_text_parts: list[str] = []
        for img in images:
            ocr_text_parts.append(_ocr_image(ocr_engine, img))
        return "\n".join(ocr_text_parts).strip(), True
    except Exception as e:
        print(f"[OCR] pdf2image OCR failed: {e}")
        return text.strip(), False


def _ocr_image(ocr_engine, img) -> str:
    """Run RapidOCR on a PIL image, return concatenated text.
    RapidOCR API: result, elapse = engine(img_or_ndarray)
    result: list of [box, text, conf] or None
    """
    import numpy as np
    arr = np.array(img)
    result, _elapse = ocr_engine(arr)
    if not result:
        return ""
    lines: list[str] = []
    for item in result:
        try:
            lines.append(item[1])
        except (IndexError, TypeError):
            continue
    return "\n".join(lines)


def _extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from .docx via python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        print(f"[OCR] docx extraction failed: {e}")
        return ""


def _extract_image_text(file_bytes: bytes) -> str:
    """OCR an image via RapidOCR."""
    ocr_engine = _get_ocr_engine()
    if ocr_engine is None:
        return ""
    try:
        from PIL import Image
        import numpy as np
        img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        return _ocr_image(ocr_engine, img)
    except Exception as e:
        print(f"[OCR] image OCR failed: {e}")
        return ""


def extract_text(filename: str, file_bytes: bytes) -> tuple[str, str, bool]:
    """Extract text from a file. Returns (text, method, used_ocr).
    method: 'pdf-electronic' | 'pdf-ocr' | 'docx' | 'image-ocr' | 'text' | 'unknown' | 'failed'
    """
    suffix = Path(filename or "").suffix.lower()
    try:
        if suffix == ".pdf":
            text, used_ocr = _extract_pdf_text(file_bytes)
            method = "pdf-ocr" if used_ocr else "pdf-electronic"
            return text, method, used_ocr
        elif suffix in (".docx",):
            return _extract_docx_text(file_bytes), "docx", False
        elif suffix in (".doc",):
            # .doc (legacy) — python-docx doesn't support; try as text, else fail
            return "", "doc-unsupported", False
        elif suffix in (".txt", ".md", ".csv", ".rtf"):
            try:
                return file_bytes.decode("utf-8", errors="ignore"), "text", False
            except Exception:
                return "", "text-failed", False
        elif suffix in (".jpg", ".jpeg", ".png", ".bmp", ".webp", ".tiff"):
            return _extract_image_text(file_bytes), "image-ocr", True
        else:
            return "", "unknown", False
    except Exception as e:
        print(f"[OCR] extract_text error: {e}")
        return "", "failed", False


# ── Local LLM structured field extraction ────────────────────
# Default: local OpenAI-compatible server (LM Studio / vLLM / llama.cpp at 127.0.0.1:8000).
# Override via env: OCR_LLM_BASE_URL / OCR_LLM_MODEL / OCR_LLM_API_KEY
# (fall back to legacy DEEPSEEK_* envs when the OCR_LLM_* ones are unset).
# Fields to extract (must match frontend form field names)
TARGET_FIELDS = [
    "applicant_name", "applicant_address", "applicant_postal_code",
    "applicant_phone", "applicant_mobile", "applicant_fax",
    "applicant_email", "applicant_other_contact",
    "respondent_name", "respondent_address", "respondent_postal_code",
    "respondent_phone", "respondent_mobile", "respondent_fax",
    "respondent_email", "respondent_other_contact",
    "mediation_willingness",  # 'mutual' | 'single_party'
    "case_facts", "dispute_matters", "mediation_demands", "demands_basis",
    "agent_name", "agent_duties",
]

EXTRACTION_PROMPT = """你是一个信息提取助手。从用户提供的调解申请书文本中，提取以下结构化字段，以 JSON 对象返回。

要求：
1. 只根据文本内容提取，未提及的字段返回空字符串 ""。
2. mediation_willingness 只能是 "mutual"（各方自愿）或 "single_party"（单方请求），无法判断则返回 ""。
3. 不要编造、不要猜测，严格基于文本。
4. 字段含义：
   - applicant_*: 申请人信息（名称/地址/邮编/固定电话/移动电话/传真/邮箱/其他联系方式）
   - respondent_*: 被申请人信息（同上）
   - mediation_willingness: 调解意愿
   - case_facts: 案件事实
   - dispute_matters: 争议事项
   - mediation_demands: 调解诉求
   - demands_basis: 理据
   - agent_name: 代理人姓名
   - agent_duties: 代理人职责
5. 返回纯 JSON，不要 markdown 代码块，不要解释。

JSON 字段结构：
{
  "applicant_name": "", "applicant_address": "", "applicant_postal_code": "",
  "applicant_phone": "", "applicant_mobile": "", "applicant_fax": "",
  "applicant_email": "", "applicant_other_contact": "",
  "respondent_name": "", "respondent_address": "", "respondent_postal_code": "",
  "respondent_phone": "", "respondent_mobile": "", "respondent_fax": "",
  "respondent_email": "", "respondent_other_contact": "",
  "mediation_willingness": "",
  "case_facts": "", "dispute_matters": "", "mediation_demands": "", "demands_basis": "",
  "agent_name": "", "agent_duties": ""
}

调解申请书文本：
"""


def extract_fields_with_llm(text: str) -> dict[str, Any]:
    """Use a local LLM (OpenAI-compatible) to extract structured fields from text. Returns dict of field -> value."""
    if not text.strip():
        return {"error": "提取文本为空"}

    # Load LLM config from env (defaults to local OpenAI-compatible server)
    api_key = os.getenv("OCR_LLM_API_KEY", os.getenv("DEEPSEEK_API_KEY", "local"))
    base_url = os.getenv("OCR_LLM_BASE_URL", os.getenv("DEEPSEEK_BASE_URL", "http://127.0.0.1:8000/v1"))
    model = os.getenv("OCR_LLM_MODEL", os.getenv("DEEPSEEK_MODEL", "Qwen3.6-35B-A3B-4bit"))

    if not api_key:
        return {"error": "未配置 LLM API Key（设置 OCR_LLM_API_KEY 或 DEEPSEEK_API_KEY），无法进行字段提取"}

    try:
        from openai import OpenAI
        client = OpenAI(api_key=api_key, base_url=base_url)

        # Truncate very long text to avoid token limits
        max_chars = 8000
        truncated = text[:max_chars] if len(text) > max_chars else text

        messages = [
            {"role": "system", "content": "你是一个严格的信息提取助手，只返回 JSON。"},
            {"role": "user", "content": EXTRACTION_PROMPT + truncated},
        ]

        try:
            response = client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=0.1,
                response_format={"type": "json_object"},
            )
        except Exception:
            # Some local servers (e.g. llama.cpp) don't support response_format; retry without it
            response = client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=0.1,
            )
        content = response.choices[0].message.content or "{}"

        import json
        fields = json.loads(content)
        # Filter to only known fields
        result: dict[str, Any] = {}
        for f in TARGET_FIELDS:
            val = fields.get(f, "")
            if isinstance(val, str):
                result[f] = val.strip()
            else:
                result[f] = str(val).strip()
        return result
    except Exception as e:
        return {"error": f"LLM 调用失败: {e}"}