"""报告生成器：输出 .docx 或 .md 文件。

反脱敏后的最终文本通过此模块持久化为报告文件。
"""

import os
from datetime import datetime

from config import settings


def save_report(case_id: str, content: str, fmt: str = "docx") -> str:
    """生成报告文件，返回绝对路径。

    case_id: 案件 ID（用于文件名）
    content: 报告内容（反脱敏后的最终文本）
    fmt: 'docx' 或 'md'
    """
    os.makedirs(settings.REPORT_DIR, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    # 清洗 case_id 中的路径分隔符，避免目录穿越
    safe_id = str(case_id).replace("/", "_").replace("\\", "_")
    filename = f"{safe_id}_{ts}.{fmt}"
    filepath = os.path.abspath(os.path.join(settings.REPORT_DIR, filename))

    if fmt == "md":
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
    elif fmt == "docx":
        from docx import Document
        doc = Document()
        for para in content.split("\n"):
            if para.strip():
                doc.add_paragraph(para)
            else:
                doc.add_paragraph("")
        doc.save(filepath)
    else:
        raise ValueError(f"不支持的格式: {fmt}（仅支持 docx 或 md）")

    return filepath
