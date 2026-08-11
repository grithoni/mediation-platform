from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = "/Users/honi/Desktop/projects/mediation-platform/.artifacts/五阶段调解方法专利申请文本-完善版.docx"

BODY_FONT = "Songti SC"
HEAD_FONT = "Heiti SC"


def set_run_font(run, name=BODY_FONT, size=Pt(11), bold=False, color=RGBColor(0, 0, 0)):
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), name)


def set_para(paragraph, before=0, after=6, line=1.15, first_line=True):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(22)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_run_font(run, HEAD_FONT, Pt(15 if level == 1 else 13), True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    set_para(p)
    set_run_font(p.add_run(text), BODY_FONT, Pt(11))
    return p


def add_center_title(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    set_run_font(run, HEAD_FONT, Pt(size), True)
    return p


def add_numbered_claim(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"{number}. {text}")
    set_run_font(r, BODY_FONT, Pt(11))


def shade_cell(cell, fill="F2F4F7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, HEAD_FONT if bold else BODY_FONT, Pt(10), bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True, "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for i, item in enumerate(row):
            set_cell_text(cells[i], item)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    return table


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


claims = [
    "一种基于五阶段结构化方法论的AI辅助商事调解方法，其特征在于，包括：接收争议案件材料、当事人陈述和调解过程记录，建立包含案件事实、主体、证据来源、争点、阶段状态和风险状态的案件状态数据结构；由阶段状态机按照接案准备阶段V、开启过程阶段A、倾听理解阶段L、方案验证阶段U和促成解决阶段E执行调解辅助流程；在每一阶段调用与该阶段绑定的AI技能节点，生成结构化阶段输出；由技能编排调度器将前一阶段的结构化输出、证据来源标签和待确认项传递至后一阶段，并依据阶段结束标准、风险阈值和人工确认结果控制阶段推进、停留或回退；其中，所述五个阶段分别配置五个AI技能节点，共形成二十五个AI技能节点，且全部自然语言生成结果均经过中立性约束引擎过滤后输出。",
    "根据权利要求1所述的方法，其特征在于，所述案件状态数据结构至少包括事实字段、立场字段、利益字段、证据字段、争点字段、程序字段、方案字段、风险字段和协议字段；每一字段均关联来源标签、置信度标签和待确认标签，以使后续阶段能够区分已确认事实、当事人陈述、系统推理结果和待人工确认信息。",
    "根据权利要求1所述的方法，其特征在于，所述接案准备阶段V包括接案评估、案件摘要、争点识别、调解准备清单和进入调解建议五个AI技能节点；其中接案评估节点基于争议类型、法律关系复杂度、信息完整度、冲突强度、授权状态和调解禁入条件生成调解适配度评分，并在评分低于预设阈值时输出补充材料任务或暂缓调解建议。",
    "根据权利要求1所述的方法，其特征在于，所述开启过程阶段A包括开场词生成、规则制定、议程设计、氛围建立和程序确认五个AI技能节点；其中程序确认节点对参与人身份、授权范围、代理权限、保密边界、录音录像许可、线上线下技术条件和单独会谈安排进行校验，并将未通过校验的事项写入所述案件状态数据结构的待确认标签。",
    "根据权利要求1所述的方法，其特征在于，所述倾听理解阶段L包括中立提问、复述确认、信息澄清、情绪降温和隐含利益识别五个AI技能节点；其中中立提问节点按照事实类、感受类、利益类、底线类和可接受选项类生成问题集合，复述确认节点分别生成事实复述文本、情绪承接文本和利益归纳文本。",
    "根据权利要求5所述的方法，其特征在于，所述隐含利益识别节点对当事人表达执行三层语义分解，将输入内容划分为事实、立场和利益，并进一步输出表层要求、深层利益、担忧点和可协商空间；所述输出与证据来源标签及置信度标签关联，以避免将未经确认的推理内容作为案件事实使用。",
    "根据权利要求5所述的方法，其特征在于，所述情绪降温节点基于文本情绪强度、冲突词触发次数、打断次数或语音情绪标签中的至少一种计算冲突强度值；当所述冲突强度值超过预设阈值时，生成按照共情回应、事实回归和需求表达三个顺序排列的降温话术，并触发阶段状态机保持在倾听理解阶段或回退至程序规则确认。",
    "根据权利要求1所述的方法，其特征在于，所述方案验证阶段U包括方案头脑风暴、方案重构、方案比较、风险分析和方案优先级五个AI技能节点；其中方案重构节点将当事人的固定立场转换为可交换条件、可组合条款和备选履行路径，并生成候选方案集合。",
    "根据权利要求8所述的方法，其特征在于，所述方案比较节点依据成本、时间、履约风险、沟通风险、后续争议风险、可执行性、关系影响和公平性生成多维评分矩阵；所述风险分析节点为每一候选方案生成风险点、触发条件、影响等级、发生概率和缓释措施；所述方案优先级节点基于所述多维评分矩阵和风险分析结果输出排序逻辑，而不直接替当事人作出最终选择。",
    "根据权利要求1所述的方法，其特征在于，所述促成解决阶段E包括决策推进、条款草拟、协议校对、履行计划和复盘总结五个AI技能节点；其中条款草拟节点将经双方确认的共识转换为包含条款标题、义务主体、行为内容、金额、期限、条件、例外、违约后果和履行验证方式的结构化条款。",
    "根据权利要求10所述的方法，其特征在于，所述协议校对节点对协议草案执行主体一致性校验、金额一致性校验、日期和期限校验、条件触发校验、例外事项校验、条款冲突校验和不可执行风险校验，并对未确认或冲突内容生成需人工确认事项。",
    "根据权利要求1所述的方法，其特征在于，所述中立性约束引擎包括材料依据校验规则、非裁判化表达规则、非替代决策规则、非偏向性语言规则和信息不足标注规则；当自然语言生成结果包含未由案件材料支持的事实、评价对错的语句、要求一方让步的指令或替代当事人决策的结论时，所述中立性约束引擎对该结果进行改写、拦截或标注。",
    "根据权利要求1所述的方法，其特征在于，所述技能编排调度器包括技能注册模块、阶段识别模块、输入完整度检测模块、技能调用模块、输出校验模块和阶段转移模块；所述技能注册模块保存二十五个AI技能节点的技能标识、所属阶段、输入规范、输出规范、触发条件和依赖关系。",
    "根据权利要求1所述的方法，其特征在于，所述阶段结束标准包括：接案准备阶段的信息完整度和调解适配度达到预设条件；开启过程阶段的身份、授权和程序事项完成确认；倾听理解阶段的关键争点、事实澄清项和双方利益覆盖度达到预设条件；方案验证阶段存在至少一个满足风险阈值和可执行性阈值的候选方案；促成解决阶段的协议条款通过一致性校验并生成履行计划。",
    "根据权利要求1所述的方法，其特征在于，当任一阶段检测到证据来源缺失、授权不足、程序确认失败、冲突强度超阈值、关键事实矛盾、候选方案风险超阈值或协议条款冲突时，阶段状态机触发回退机制，并将回退原因、需补充信息和目标阶段写入案件状态数据结构。",
    "一种基于五阶段结构化方法论的AI辅助商事调解系统，其特征在于，包括：输入模块，用于接收争议案件材料、当事人陈述和调解过程记录；案件状态管理模块，用于建立并维护包含事实、立场、利益、证据、争点、程序、方案、风险和协议的案件状态数据结构；阶段状态机模块，用于按照接案准备阶段V、开启过程阶段A、倾听理解阶段L、方案验证阶段U和促成解决阶段E控制调解辅助流程；技能编排调度模块，用于注册、调用和校验二十五个AI技能节点；中立性约束模块，用于对自然语言生成结果执行中立性过滤；隐含利益推理模块，用于生成表层要求、深层利益、担忧点和可协商空间；方案闭环评估模块，用于生成候选方案集合、多维评分矩阵、风险分析结果和方案优先级；协议生成校对模块，用于生成结构化条款、协议校对结果和履行计划；输出模块，用于输出各阶段调解辅助文档。",
    "根据权利要求16所述的系统，其特征在于，所述技能编排调度模块包括技能注册单元、依赖关系管理单元、阶段输入检测单元、技能调用单元、输出格式校验单元和阶段转移判断单元。",
    "根据权利要求16所述的系统，其特征在于，所述中立性约束模块包括材料依据校验单元、偏向性表达检测单元、裁判化表述检测单元、替代决策检测单元和待确认标注单元。",
    "根据权利要求16所述的系统，其特征在于，所述系统以SaaS平台、调解终端设备、移动端应用或私有化部署服务器中的任一种形式实现。",
    "一种电子设备，包括处理器和存储器，所述存储器中存储有计算机程序，所述计算机程序被所述处理器执行时实现权利要求1至15任一项所述的方法。",
    "一种计算机可读存储介质，其上存储有计算机程序，所述计算机程序被处理器执行时实现权利要求1至15任一项所述的方法。",
]


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
styles["Normal"].font.name = BODY_FONT
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
styles["Normal"].font.size = Pt(11)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.15

for style_name, size in (("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12)):
    style = styles[style_name]
    style.font.name = HEAD_FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), HEAD_FONT)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(6)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(footer.add_run("第 "), BODY_FONT, Pt(9))
add_page_number(footer)
set_run_font(footer.add_run(" 页"), BODY_FONT, Pt(9))

add_center_title(doc, "专利申请文件", 18)
add_heading(doc, "发明名称", 1)
add_body(doc, "基于五阶段结构化方法论的AI辅助商事调解方法及系统")

add_heading(doc, "摘要", 1)
add_body(
    doc,
    "本发明公开了一种基于五阶段结构化方法论的AI辅助商事调解方法及系统。该方法接收争议案件材料、"
    "当事人陈述和调解过程记录，建立包含事实、立场、利益、证据、争点、程序、方案、风险和协议的案件状态"
    "数据结构；通过阶段状态机依次执行接案准备V、开启过程A、倾听理解L、方案验证U和促成解决E五个阶段，"
    "并在每一阶段调用对应的五个AI技能节点，共二十五个AI技能节点。系统通过技能编排调度器实现阶段间信息"
    "流转、输入完整度检测、输出校验和阶段回退；通过中立性约束引擎对自然语言生成结果进行材料依据校验、"
    "非裁判化表达校验和非替代决策校验；通过隐含利益推理模块区分事实、立场和利益，并输出表层要求、深层利益、"
    "担忧点和可协商空间；通过方案闭环评估模块完成方案重构、多维评分、风险分析和优先级排序；通过协议生成"
    "校对模块生成结构化条款、校对冲突和履行计划。该方案使调解辅助过程具备标准化、可追踪、可审计和可执行"
    "的技术效果。"
)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
add_center_title(doc, "权利要求书", 16)
for i, claim in enumerate(claims, 1):
    add_numbered_claim(doc, i, claim)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
add_center_title(doc, "说明书", 16)
add_heading(doc, "技术领域", 1)
add_body(doc, "本发明涉及人工智能辅助争议解决和自然语言信息处理技术领域，具体涉及一种基于五阶段结构化方法论的AI辅助商事调解方法及系统。")

add_heading(doc, "背景技术", 1)
for text in [
    "商事调解作为一种高效、灵活的争议解决方式，在合同争议、交易纠纷、平台纠纷和企业间纠纷处理中具有重要作用。传统调解高度依赖调解员个人经验，容易出现流程边界不清、信息记录离散、争点识别不稳定、方案比较缺少统一维度以及协议条款难以落地等问题。",
    "现有AI调解辅助技术通常集中在情绪识别、话术推荐、协议生成或案例检索等单点能力。此类技术虽然能够提升某一环节的效率，但并未将调解活动抽象为可被机器执行和审计的阶段状态机，也缺少跨阶段数据结构、技能编排、证据来源标签和中立性约束机制。",
    "例如，部分方案通过情绪评分和策略推荐辅助调解，但缺少五阶段流程的结束标准和回退机制；部分方案能够生成协议文本，但未将协议文本与前序争点、利益识别、风险分析和人工确认项关联；还有部分方案提供话术推荐，却没有对生成内容进行材料依据校验和非替代决策约束。",
    "因此，现有技术仍存在以下不足：第一，调解流程未被建模为具有输入、输出、结束标准和回退条件的阶段状态机；第二，AI技能以单点方式工作，缺少二十五个技能节点之间的依赖关系和调度机制；第三，自然语言生成容易产生偏向性、裁判化或超出材料依据的内容；第四，隐含利益推理缺少事实、立场、利益的分层处理；第五，候选方案生成、风险评估、协议校对和履行计划未形成闭环。"
]:
    add_body(doc, text)

add_heading(doc, "发明内容", 1)
add_heading(doc, "要解决的技术问题", 2)
add_body(doc, "本发明旨在解决现有AI调解辅助技术中流程不可标准化、技能不可编排、输出不可审计、中立性不可控制、方案评估不闭环以及协议执行风险难以前置识别的技术问题。")

add_heading(doc, "技术方案", 2)
for text in [
    "本发明采用V-A-L-U-E五阶段结构化方法论框架，将商事调解全流程划分为接案准备、开启过程、倾听理解、方案验证和促成解决五个阶段。系统为每个阶段配置五个AI技能节点，共二十五个AI技能节点，并通过阶段状态机和技能编排调度器控制各节点的调用顺序、依赖关系、输入完整度和输出合规性。",
    "区别于仅进行话术生成或协议生成的单点技术，本发明建立案件状态数据结构，并在其中维护事实、立场、利益、证据、争点、程序、方案、风险和协议字段。每个字段均关联来源标签、置信度标签和待确认标签，使系统能够在后续阶段区分材料事实、当事人陈述、系统推理和待人工确认信息。",
    "系统进一步设置中立性约束引擎，在自然语言生成前后执行材料依据校验、非裁判化表达校验、非替代决策校验、非偏向性语言校验和信息不足标注。当生成结果超出材料依据、评价对错、替代当事人作出决定或要求一方让步时，系统对结果进行改写、拦截或标注。",
    "在倾听理解阶段，系统通过隐含利益推理模块对当事人表达执行三层语义分解，区分事实、立场和利益，并生成表层要求、深层利益、担忧点和可协商空间。该结果并不直接作为事实使用，而是绑定置信度和来源标签后进入方案验证阶段。",
    "在方案验证阶段，系统通过方案头脑风暴、方案重构、方案比较、风险分析和方案优先级形成闭环。候选方案被转换为可交换条件、可组合条款和备选履行路径，并通过成本、时间、履约风险、沟通风险、后续争议风险、可执行性、关系影响和公平性形成多维评分矩阵。",
    "在促成解决阶段，系统将双方确认的共识转化为结构化条款，并对主体、金额、期限、条件、例外、违约后果和履行验证方式进行一致性校验，最后生成履行计划和复盘总结。"
]:
    add_body(doc, text)

add_heading(doc, "阶段与技能节点", 2)
add_table(
    doc,
    ["阶段", "代号", "目标", "AI技能节点"],
    [
        ["接案准备", "V", "判断适配性、整理案件、识别争点", "接案评估、案件摘要、争点识别、调解准备清单、进入调解建议"],
        ["开启过程", "A", "建立安全有序的对话基础", "开场词生成、规则制定、议程设计、氛围建立、程序确认"],
        ["倾听理解", "L", "识别事实、立场、利益和真实需要", "中立提问、复述确认、信息澄清、情绪降温、隐含利益识别"],
        ["方案验证", "U", "生成、重构、比较和排序方案", "方案头脑风暴、方案重构、方案比较、风险分析、方案优先级"],
        ["促成解决", "E", "形成可执行协议和履行计划", "决策推进、条款草拟、协议校对、履行计划、复盘总结"],
    ],
)

add_heading(doc, "有益效果", 2)
for text in [
    "本发明通过阶段状态机将调解流程转化为可执行、可回退、可审计的信息处理过程，提高调解质量的一致性。",
    "本发明通过二十五个AI技能节点和技能编排调度器实现跨阶段的端到端辅助，克服单点AI能力无法覆盖完整调解流程的问题。",
    "本发明通过中立性约束引擎降低生成式AI在调解场景中的偏向性、裁判化和越权决策风险。",
    "本发明通过事实、立场、利益三层语义分解和来源标签机制，提高隐含利益推理的可解释性和可追踪性。",
    "本发明通过方案闭环评估和协议一致性校验，提高候选方案的可比较性、风险可控性和协议可执行性。",
]:
    add_body(doc, text)

add_heading(doc, "附图说明", 1)
for text in [
    "图1为本发明五阶段结构化调解流程示意图。",
    "图2为本发明案件状态数据结构示意图。",
    "图3为本发明二十五个AI技能节点及技能编排调度架构示意图。",
    "图4为本发明中立性约束引擎处理流程示意图。",
    "图5为本发明隐含利益推理流程示意图。",
    "图6为本发明方案生成与风险评估闭环示意图。",
    "图7为本发明协议生成校对与履行计划生成流程示意图。",
    "图8为本发明系统模块结构示意图。",
]:
    add_body(doc, text)

add_heading(doc, "具体实施方式", 1)
add_heading(doc, "实施例一：SaaS平台实现", 2)
for text in [
    "本实施例中，系统部署于云端服务器，用户通过Web界面上传案件申请书、合同、证据、沟通记录和当事人信息。输入模块对上传材料进行解析，并由案件状态管理模块建立案件状态数据结构。",
    "在接案准备阶段V，系统调用接案评估、案件摘要、争点识别、调解准备清单和进入调解建议五个节点。接案评估节点依据争议类型、法律关系复杂度、信息完整度、冲突强度、授权状态和调解禁入条件生成调解适配度评分；争点识别节点将争点区分为事实争点、法律争点和程序争点，并标注是否可调解。",
    "在开启过程阶段A，系统调用开场词生成、规则制定、议程设计、氛围建立和程序确认五个节点。程序确认节点对身份、授权、代理权限、保密边界、录音录像许可和单独会谈安排进行校验。若授权范围未确认，阶段状态机保持在该阶段并输出补充确认任务。",
    "在倾听理解阶段L，系统接收双方陈述、语音转写文本或调解记录。中立提问节点生成事实类、感受类、利益类、底线类和可接受选项类问题；隐含利益识别节点执行事实、立场和利益三层分解，并输出表层要求、深层利益、担忧点和可协商空间。",
    "在方案验证阶段U，系统基于双方利益和客观约束生成候选方案集合。方案重构节点将固定立场转换为可交换条件、可组合条款和备选履行路径；方案比较节点生成多维评分矩阵；风险分析节点输出触发条件、影响等级、发生概率和缓释措施；方案优先级节点仅输出排序逻辑，不替代当事人决策。",
    "在促成解决阶段E，系统将经双方确认的共识转化为结构化条款，协议校对节点对主体、金额、期限、条件、例外、违约后果和履行验证方式进行一致性校验。校验通过后，履行计划节点生成时间表、责任人、交付物、检查点和提醒机制。",
]:
    add_body(doc, text)

add_heading(doc, "实施例二：调解终端设备实现", 2)
for text in [
    "本实施例中，系统以软硬件结合的调解终端设备实现。所述终端设备包括处理器、存储器、显示屏、麦克风阵列、扬声器、摄像头和通信模块。",
    "麦克风阵列采集现场语音，摄像头采集表情和肢体动作，系统通过语音识别和情感分析生成文本和情绪标签。情绪降温节点结合文本情绪强度、打断次数和语音情绪标签计算冲突强度值，当冲突强度值超过阈值时，向调解员显示按照共情回应、事实回归和需求表达顺序排列的降温建议。",
]:
    add_body(doc, text)

add_heading(doc, "实施例三：移动端应用实现", 2)
for text in [
    "本实施例中，系统以移动端应用形式实现。调解员可手动选择当前阶段，也可由阶段状态机根据阶段结束标准自动判断是否进入下一阶段。",
    "移动端应用支持离线缓存案件材料和阶段输出；在网络恢复后，将缓存内容同步至服务器，并保留版本记录、人工确认记录和阶段回退记录，以便审计和复盘。",
]:
    add_body(doc, text)

add_heading(doc, "创造性说明", 1)
for text in [
    "本发明的创造性并非仅在于将AI用于调解，而在于将调解全过程建模为五阶段状态机，并为每一阶段配置明确的输入、输出、结束标准、回退条件和技能节点依赖关系。",
    "本发明进一步将中立性原则转化为可执行的技术约束规则，使自然语言生成结果能够被校验、改写、拦截或标注，解决了生成式AI在调解场景中可能产生偏向性和替代决策的问题。",
    "本发明还将隐含利益推理结果与来源标签、置信度标签和待确认标签绑定，使系统能够在方案生成时使用深层利益信息，同时避免将推理内容误作为已确认事实。",
    "本发明通过候选方案生成、立场重构、多维评分、风险分析、优先级排序、协议校对和履行计划形成闭环，使前序沟通信息能够转化为可比较、可验证和可执行的调解成果。",
]:
    add_body(doc, text)

doc.save(OUTPUT_PATH)
print(OUTPUT_PATH)
