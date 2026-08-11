from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = "/Users/honi/Desktop/projects/mediation-platform/.artifacts/five_phase_mediation_patent_application.docx"


ABSTRACT = (
    "本发明涉及争议解决与智能信息处理技术领域，公开了一种基于五阶段流程的调解辅助方法、"
    "系统、设备及存储介质。所述方法包括：获取争议案件材料并执行接案准备阶段处理，形成适配度"
    "评估结果、案件摘要、争点清单和调解准备清单；执行开启过程阶段处理，生成开场指引、会议规则、"
    "议程方案和程序确认结果；执行倾听理解阶段处理，基于双方陈述生成中立提问、复述确认、信息澄清项、"
    "情绪降温建议及隐含利益识别结果；执行方案验证阶段处理，形成候选方案集合、方案比较结果、风险分析"
    "结果及优先级排序结果；执行促成解决阶段处理，生成决策推进指引、和解条款草案、协议校对结果、履行"
    "计划及复盘结果；并依据阶段转移规则驱动各阶段有序切换和结果回写。该方案实现了调解流程结构化、"
    "输出标准化、决策过程可追踪以及协议结果可执行性提升。"
)


CLAIMS = [
    "一种基于五阶段流程的调解辅助方法，其特征在于，包括：获取争议案件材料；"
    "对所述争议案件材料执行接案准备处理，以输出争议适配度评估结果、案件摘要、争点识别结果、"
    "调解准备清单及进入调解决策建议；在满足预设进入条件后执行开启过程处理，以输出开场词、会议规则、"
    "阶段议程、氛围建立引导语及程序确认结果；基于调解过程中的交互内容执行倾听理解处理，以输出中立提问"
    "集合、复述确认结果、信息澄清项、情绪降温响应及隐含利益识别结果；基于所述隐含利益识别结果和争议"
    "约束条件执行方案验证处理，以输出候选方案集合、方案比较结果、风险分析结果及方案优先级结果；基于"
    "优先级结果和当事人确认信息执行促成解决处理，以输出决策推进问题、和解条款草案、协议校对结果、"
    "履行计划及复盘总结；其中，系统依据阶段转移规则控制五阶段顺序推进、回退或停留，并将各阶段输出"
    "结构化存储。",
    "根据权利要求1所述的方法，其特征在于，所述接案准备处理包括：对案件材料进行事实抽取、"
    "主体识别、时间线整理、证据归类和争点分类，并基于争议类型、冲突程度、信息完整度及调解前置条件"
    "生成调解适配度评分。",
    "根据权利要求1所述的方法，其特征在于，所述开启过程处理包括：根据参与人身份、授权范围、"
    "语言安排、线上或线下场景以及是否允许单独会谈，自动生成会议规则、议程时长分配和程序确认清单。",
    "根据权利要求1所述的方法，其特征在于，所述倾听理解处理包括：对当事人陈述进行事实语句、"
    "情绪语句和利益语句分层识别，并分别生成事实复述文本、情绪承接文本和利益归纳文本。",
    "根据权利要求1所述的方法，其特征在于，所述方案验证处理包括：对各方立场进行可交换条件重构，"
    "将对立立场转换为多个可组合选项，并按照成本、时间、风险、可执行性、关系影响及公平性维度对候选"
    "方案进行比较排序。",
    "根据权利要求1所述的方法，其特征在于，所述促成解决处理包括：将已确认共识转换为包含条款标题、"
    "条款正文、履行期限、责任主体和待补全占位字段的结构化协议草案，并对日期、金额、条件触发项、"
    "例外事项和执行方式进行一致性校验。",
    "根据权利要求1所述的方法，其特征在于，所述阶段转移规则至少包括：接案准备阶段输出的信息"
    "完整度达到阈值时进入开启过程阶段；程序确认完成时进入倾听理解阶段；争议关键信息澄清完成且双方"
    "利益识别达到预设覆盖度时进入方案验证阶段；候选方案形成且存在可接受方案区间时进入促成解决阶段。",
    "根据权利要求1所述的方法，其特征在于，当任一阶段检测到关键信息缺失、授权不足、情绪冲突"
    "超阈值或协议内容冲突时，触发回退机制，返回前序阶段补充处理。",
    "一种基于五阶段流程的调解辅助系统，其特征在于，包括：案件接入模块，用于获取争议案件材料；"
    "接案准备模块，用于输出适配度评估结果、案件摘要、争点清单和准备清单；开启过程模块，用于输出开场词、"
    "会议规则、议程方案和程序确认结果；倾听理解模块，用于输出提问结果、复述结果、澄清结果、降温结果"
    "及利益识别结果；方案验证模块，用于输出候选方案、比较结果、风险结果及优先级结果；促成解决模块，"
    "用于输出条款草案、校对结果、履行计划和复盘结果；流程编排模块，用于依据阶段转移规则控制各模块之间"
    "的调用顺序；存储模块，用于存储案件数据、阶段结果和版本记录。",
    "一种电子设备，包括处理器和存储器，所述存储器中存储有程序，所述程序被所述处理器执行时实现"
    "权利要求1至8任一项所述的方法。",
    "一种计算机可读存储介质，其上存储有计算机程序，所述计算机程序被处理器执行时实现权利要求1至8"
    "任一项所述的方法。",
]


DESCRIPTION_SECTIONS = [
    (
        "一、技术领域",
        [
            "本发明涉及争议解决信息处理技术领域，尤其涉及一种基于五阶段流程的调解辅助方法、系统、设备及存储介质。"
        ],
    ),
    (
        "二、背景技术",
        [
            "现有调解活动较多依赖调解员个人经验，存在流程不统一、阶段边界不清、关键信息提取不稳定、"
            "候选方案生成缺乏结构化比较、和解条款可执行性不足等问题。",
            "尤其在商事纠纷、平台调解和在线调解场景中，案件材料来源多样、参与主体复杂、沟通轮次较多，"
            "若缺乏统一的流程编排与结构化输出机制，容易导致前期评估不充分、中期沟通失焦、后期协议遗漏关键"
            "履行条件，影响调解效率与结果质量。",
            "因此，需要一种能够对调解全过程进行阶段化组织、规则化推进、结构化输出和结果可追踪管理的技术方案。"
        ],
    ),
    (
        "三、发明内容",
        [
            "本发明的目的在于提供一种基于五阶段流程的调解辅助方法、系统、设备及存储介质，以解决现有技术中"
            "调解流程标准化不足、信息处理离散、方案比较不充分以及协议执行性弱的问题。",
            "为实现上述目的，本发明采用如下技术方案：将调解过程划分为接案准备、开启过程、倾听理解、方案验证"
            "和促成解决五个阶段；针对每一阶段分别设置输入数据结构、处理规则、输出模板和阶段转移条件；通过"
            "流程编排模块控制阶段推进、停留或回退；并对全过程中产生的评估结果、引导内容、候选方案及协议文本"
            "进行结构化存储。",
            "与现有技术相比，本发明至少具有如下有益效果：",
        ],
    ),
    (
        "四、附图说明",
        [
            "图1为本发明方法的整体流程示意图。",
            "图2为接案准备阶段的数据处理流程示意图。",
            "图3为开启过程阶段的程序确认与议程生成流程示意图。",
            "图4为倾听理解阶段的陈述分层处理示意图。",
            "图5为方案验证阶段的候选方案生成与排序流程示意图。",
            "图6为促成解决阶段的条款生成与协议校对流程示意图。",
            "图7为本发明系统的模块结构示意图。",
        ],
    ),
    (
        "五、具体实施方式",
        [
            "在一实施例中，系统首先接收案件申请书、合同文本、证据材料、沟通记录和当事人基础信息，并对输入内容"
            "执行接案准备处理。该阶段中，系统抽取案件主体、争议事实、关键时间节点及证据目录，形成一页式案件摘要；"
            "同时识别事实争点、法律争点和程序争点，并结合信息完整度、冲突程度和前置障碍生成调解适配度。若存在"
            "授权不明、材料缺失或不宜调解情形，则输出补充建议或暂缓进入结果。",
            "在进入开启过程阶段后，系统根据调解形式、参与人身份及授权状态，自动生成开场词、会议规则、议程安排、"
            "暖场引导语以及程序确认清单。程序确认清单至少覆盖身份核验、代理权限、语言安排、技术平台、录音录像"
            "许可及是否进行单独会谈等事项。",
            "在倾听理解阶段，系统接收双方陈述文本、语音转写内容或调解记录，将表达内容划分为事实信息、情绪表达和"
            "利益诉求三类；针对事实信息输出中立提问和事实复述文本，针对情绪表达输出承接与降温文本，针对利益诉求"
            "输出利益归纳和可协商空间识别结果；同时对存在矛盾、模糊或缺证的信息生成澄清问题及所需证据项。",
            "在方案验证阶段，系统基于识别出的利益诉求、争点边界和客观约束条件生成多个候选方案方向，并将刚性立场"
            "重构为可交换条件和可组合条款；随后按照成本、时间、风险、可执行性、关系影响和公平性维度对方案进行"
            "比较，形成排序建议；若发现关键风险未被覆盖，则要求补充验证信息或回退至倾听理解阶段。",
            "在促成解决阶段，系统依据已确认的共识生成结构化和解条款草案，条款内容至少包括付款义务、行为义务、"
            "时间节点、违约处理、保密要求和争议后续安排中的一项或多项；之后执行协议校对，检查金额、日期、条件"
            "触发、例外情形、履行方式和责任主体是否存在歧义、遗漏或冲突；校对完成后进一步生成履行计划和复盘总结。",
            "在另一实施例中，流程编排模块维护每一阶段的状态标识、输入完整度、风险等级和转移条件。当某一阶段的"
            "关键指标未达到预设阈值时，系统不进入下一阶段，而是返回缺口清单并触发补充处理，从而实现全过程闭环控制。",
            "在又一实施例中，本发明可以部署于调解平台服务器、本地终端或云端协同系统中，由处理器执行程序指令实现"
            "上述方法步骤，各阶段输出均以结构化字段形式保存，便于检索、审计、复盘和后续统计分析。"
        ],
    ),
]


BENEFITS = [
    "将调解全过程拆解为统一的五阶段流程，降低对个人经验的依赖。",
    "通过阶段化输出模板提升案件摘要、争点识别、方案比较和协议草拟的一致性。",
    "通过转移规则和回退机制实现过程可追踪与风险可控。",
    "通过协议一致性校验和履行计划生成，提高和解结果的明确性和可执行性。",
    "适用于线下调解、在线调解、平台调解及人机协同调解场景。",
]


def set_font(run, name="STSong", size=Pt(12), bold=False, color=RGBColor(0, 0, 0)):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color


def set_paragraph_format(paragraph, before=0, after=6, line=1.15, first_line_chars=2):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line_chars is not None:
        fmt.first_line_indent = Pt(24)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
styles["Normal"].font.name = "STSong"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "STSong")
styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "STSong")
styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "STSong")
styles["Normal"]._element.rPr.rFonts.set(qn("w:cs"), "STSong")
styles["Normal"].font.size = Pt(12)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.15

for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
    style = styles[style_name]
    style.font.name = "Heiti SC"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti SC")
    style._element.rPr.rFonts.set(qn("w:ascii"), "Heiti SC")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Heiti SC")
    style._element.rPr.rFonts.set(qn("w:cs"), "Heiti SC")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_before = Pt(14 if style_name == "Heading 1" else 10)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

footer_para = section.footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_para.add_run("第 ")
set_font(run, size=Pt(10))
add_page_number(footer_para)
run = footer_para.add_run(" 页")
set_font(run, size=Pt(10))

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(18)
title_run = title.add_run("一种基于五阶段流程的调解辅助方法、系统、设备及存储介质")
set_font(title_run, name="Heiti SC", size=Pt(18), bold=True)

for label in ["摘要", "权利要求书", "说明书"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(label)
    set_font(r, name="Heiti SC", size=Pt(16), bold=True)

    if label == "摘要":
        body = doc.add_paragraph()
        set_paragraph_format(body)
        set_font(body.add_run(ABSTRACT))
        doc.add_paragraph()
    elif label == "权利要求书":
        for index, claim in enumerate(CLAIMS, start=1):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Pt(0)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            set_font(p.add_run(claim))
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    else:
        for section_title, paragraphs in DESCRIPTION_SECTIONS:
            h = doc.add_paragraph(style="Heading 1")
            h_run = h.add_run(section_title)
            set_font(h_run, name="Heiti SC", size=Pt(14), bold=True)
            for text in paragraphs:
                p = doc.add_paragraph()
                set_paragraph_format(p)
                set_font(p.add_run(text))
            if section_title == "三、发明内容":
                for item in BENEFITS:
                    bp = doc.add_paragraph(style="List Bullet")
                    bp.paragraph_format.space_after = Pt(6)
                    bp.paragraph_format.line_spacing = 1.15
                    set_font(bp.add_run(item))

doc.save(OUTPUT_PATH)
print(OUTPUT_PATH)
