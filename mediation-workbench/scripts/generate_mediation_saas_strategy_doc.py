from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_TblPr
from docx.oxml.text.paragraph import CT_P
from docx.shared import Cm, Inches, Pt, RGBColor


TITLE = "借鉴 Harvey 模式优化调解平台的架构体系、功能与商业化建议"
SUBTITLE = "定位：面向调解机构内部提效的工作台 SaaS"


@dataclass(frozen=True)
class Tokens:
    body_font: str = "Hiragino Sans GB"
    body_size: int = 11
    h1_size: int = 16
    h2_size: int = 13
    h3_size: int = 12
    title_size: int = 22
    subtitle_size: int = 11
    blue: str = "2E74B5"
    dark_blue: str = "1F4D78"
    ink: str = "1F2937"
    muted: str = "6B7280"
    table_fill: str = "F2F4F7"
    callout_fill: str = "F7FAFC"


TOKENS = Tokens()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def remove_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = CT_TblPr.new_tblPr()
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "D9E2F0")
        borders.append(elem)
    tbl_pr.append(borders)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def style_run(run, *, size: int, bold=False, color: str = TOKENS.ink, italic=False) -> None:
    font = run.font
    font.name = TOKENS.body_font
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic
    font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:ascii"), TOKENS.body_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), TOKENS.body_font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), TOKENS.body_font)
    run._element.rPr.rFonts.set(qn("w:cs"), TOKENS.body_font)


def set_paragraph_style(paragraph, *, before=0, after=6, line=1.1, alignment=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    paragraph.alignment = alignment


def add_body_paragraph(doc: Document, text: str, *, before=0, after=6) -> None:
    p = doc.add_paragraph()
    set_paragraph_style(p, before=before, after=after)
    run = p.add_run(text)
    style_run(run, size=TOKENS.body_size)


def add_lead_paragraph(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_style(p, before=0, after=8)
    label_run = p.add_run(f"{label}：")
    style_run(label_run, size=TOKENS.body_size, bold=True, color=TOKENS.dark_blue)
    text_run = p.add_run(text)
    style_run(text_run, size=TOKENS.body_size)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    if level == 1:
        set_paragraph_style(p, before=16, after=8)
        size = TOKENS.h1_size
        color = TOKENS.blue
    elif level == 2:
        set_paragraph_style(p, before=12, after=6)
        size = TOKENS.h2_size
        color = TOKENS.blue
    else:
        set_paragraph_style(p, before=8, after=4)
        size = TOKENS.h3_size
        color = TOKENS.dark_blue
    run = p.add_run(text)
    style_run(run, size=size, bold=True, color=color)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(6.5)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, TOKENS.callout_fill)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)

    p1 = cell.paragraphs[0]
    set_paragraph_style(p1, before=0, after=4)
    r1 = p1.add_run(title)
    style_run(r1, size=11, bold=True, color=TOKENS.dark_blue)

    p2 = cell.add_paragraph()
    set_paragraph_style(p2, before=0, after=0)
    r2 = p2.add_run(body)
    style_run(r2, size=11)

    spacer = doc.add_paragraph()
    set_paragraph_style(spacer, before=0, after=6)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.allow_autofit = False
    remove_table_borders(table)

    for index, width in enumerate(widths):
        table.columns[index].width = Inches(width)

    hdr_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        cell = hdr_cells[index]
        cell.width = Inches(widths[index])
        set_cell_margins(cell)
        set_cell_shading(cell, TOKENS.table_fill)
        p = cell.paragraphs[0]
        set_paragraph_style(p, before=0, after=0)
        r = p.add_run(header)
        style_run(r, size=10, bold=True, color=TOKENS.dark_blue)

    for row_values in rows:
        row_cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cell = row_cells[index]
            cell.width = Inches(widths[index])
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            set_paragraph_style(p, before=0, after=0)
            r = p.add_run(value)
            style_run(r, size=10)

    doc.add_paragraph()


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_style(p, before=0, after=0, line=1.0, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    r = p.add_run("调解平台内部战略建议 | 第 ")
    style_run(r, size=9, color=TOKENS.muted)
    add_page_number(p)
    r2 = p.add_run(" 页")
    style_run(r2, size=9, color=TOKENS.muted)


def build_doc(output_path: Path) -> None:
    doc = Document()
    configure_page(doc)

    title = doc.add_paragraph()
    set_paragraph_style(title, before=0, after=6, line=1.0, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    r = title.add_run(TITLE)
    style_run(r, size=TOKENS.title_size, bold=True, color=TOKENS.dark_blue)

    subtitle = doc.add_paragraph()
    set_paragraph_style(subtitle, before=0, after=12, line=1.0)
    sr = subtitle.add_run(SUBTITLE)
    style_run(sr, size=TOKENS.subtitle_size, color=TOKENS.muted, italic=True)

    add_callout(
        doc,
        "核心判断",
        "如果借鉴 Harvey 的演进路径，这个项目最值得投入的方向不是继续堆单点 AI 功能，而是把现有工作台、知识库、案件分析和协同能力，升级为一套面向调解机构内部的“调解操作系统”。系统应围绕案件全生命周期，提供任务编排、知识调用、角色协作、结果留痕和组织治理能力，让 AI 从问答助手转变为可控、可审计、可复用的专业代理。",
    )

    add_heading(doc, "一、当前项目与 Harvey 模式的映射判断", 1)
    add_lead_paragraph(doc, "现状优势", "项目已经具备官网引流、当事人建案、调解员工作台、知识库检索、OCR、案件分析、VALUE 能力和小程序接口，基础形态接近 Harvey 早期的“工作台 + 知识 + AI 分析”组合。")
    add_lead_paragraph(doc, "主要短板", "目前能力仍然偏功能堆叠，AI 主要以分析接口和聊天能力存在，尚未升级为可编排的案件工作流引擎，也缺少组织级知识沉淀、权限治理、标准化模板和复盘闭环。")
    add_lead_paragraph(doc, "Harvey 可借鉴之处", "Harvey 的关键不是模型更强，而是把复杂专业工作拆成 Plan、Research、Work、Deliver、Review 的代理闭环，并把知识、协作、安全与审计整合进同一平台。调解场景同样适合这种范式。")

    add_heading(doc, "二、面向调解机构内部提效的目标产品定义", 1)
    add_body_paragraph(doc, "建议把平台重新定义为“调解机构工作台 SaaS”，而不是“通用在线调解网站”。前者强调机构生产力、案件周转效率、标准化交付、知识复用和管理可见性，商业上更容易形成持续付费、实施服务和续约逻辑。")
    add_table(
        doc,
        ["维度", "当前形态", "建议升级形态", "业务价值"],
        [
            ["产品定位", "业务系统 + AI 功能", "调解机构工作台 SaaS", "从项目制交付走向订阅制收入"],
            ["AI 角色", "问答与分析辅助", "案件执行代理与副驾驶", "提高单个调解员人均产能"],
            ["知识管理", "零散检索与材料查看", "机构知识底座 + 模板沉淀", "形成组织复利"],
            ["协作模式", "单案查看与处理", "角色分工 + 留痕协同", "提升案件流转效率"],
            ["管理能力", "个体操作视角", "机构级看板与治理", "支持负责人管控质量和产能"],
        ],
        [1.1, 1.4, 2.2, 1.8],
    )

    add_heading(doc, "三、建议的目标技术架构体系", 1)
    add_body_paragraph(doc, "建议把当前 Nuxt 单体工作台逐步演进为“模块化单体 + 后台任务运行时”的结构，而不是一开始拆成大量微服务。第一阶段继续以 `mediation-workbench` 为主应用壳，内部按领域模块拆边界；第二阶段再把高负载和高异步能力拆成独立进程。这样既能控制复杂度，也方便保留现有 SQLite / Drizzle / H3 体系。")
    add_table(
        doc,
        ["架构层", "建议组件", "主要职责", "技术重点"],
        [
            ["接入层", "Nuxt 页面、公开 API、小程序 API、WebSocket", "统一承载当事人端、调解员端、开放接口与实时消息", "BFF 聚合、鉴权、限流、会话隔离"],
            ["领域层", "案件域、调解流程域、文书域、知识域、组织治理域", "封装核心业务规则与状态机", "模块边界、领域事件、幂等更新"],
            ["代理层", "Agent Runtime、Prompt 模板、工具适配器、任务队列", "执行案件分诊、分析、文书生成、复盘沉淀", "异步任务、重试、超时、人工接管"],
            ["知识层", "OCR、Chunking、Embedding、检索索引、引用服务", "完成材料解析、知识召回、证据定位与出处追踪", "混合检索、版本化索引、权限过滤"],
            ["治理层", "RBAC、审计日志、配置中心、监控告警", "保证多机构 SaaS 的权限、可观测性和合规", "租户隔离、追踪链路、配置灰度"],
        ],
        [1.0, 1.7, 2.0, 1.8],
    )
    add_heading(doc, "1. 接入层与 BFF", 2)
    add_body_paragraph(doc, "建议把 `pages/party`、`pages/mediator`、`server/api` 和 `server/mp` 看作多个前端/渠道入口，上层统一接入 BFF 逻辑。BFF 的职责不是承载复杂业务，而是负责会话、鉴权、参数标准化、聚合视图和渠道差异适配。这样可以避免把案件规则继续散落在页面逻辑和零散 API 中。")
    add_heading(doc, "2. 领域服务层", 2)
    add_body_paragraph(doc, "建议至少拆成五个内部领域模块：`case-domain`、`workflow-domain`、`document-domain`、`knowledge-domain`、`org-domain`。每个模块输出稳定的 service 接口和 repository 接口，禁止页面直接拼 SQL 或跨模块读写。以案件状态迁移、代理结果落库、审批结论写回为例，都应先经过领域服务封装。")
    add_heading(doc, "3. 代理执行运行时", 2)
    add_body_paragraph(doc, "参考 Harvey 的 Plan、Research、Work、Deliver、Review 闭环，建议在服务端新增 Agent Runtime。Runtime 需要有任务表、步骤表、输入输出快照、工具调用记录、重试策略、超时控制和人工接管标记。这样 AI 不再只是一次性接口调用，而是可恢复、可审计、可中断的长任务执行。")
    add_heading(doc, "4. 知识检索与证据引用层", 2)
    add_body_paragraph(doc, "现有 KB 与 OCR 应升级为可版本化的知识流水线：文档上传后先做 OCR / 清洗 / 分段 / 向量化 / 元数据标注，再写入检索索引。检索时不仅返回答案，还要返回 chunk 来源、页码、原文片段和权限校验结果，使代理输出具备可引用、可追责的证据链。")
    add_heading(doc, "5. 治理与租户隔离层", 2)
    add_body_paragraph(doc, "内部提效型 SaaS 最终会面对多机构、多团队、多角色共存。建议从现在开始按租户维度设计数据模型与权限模型，即使底层仍是单库，也要在表结构、查询条件和日志上下文中保留 `tenant_id`、`workspace_id`、`actor_id` 等字段，为后续 SaaS 化和私有化部署做准备。")

    add_heading(doc, "四、优先上线的核心功能包", 1)
    add_heading(doc, "1. 案件分诊代理", 2)
    add_body_paragraph(doc, "系统自动识别案件类型、争议焦点、证据完整度、紧急程度和建议流转路径，帮助前台受理或案管人员快速完成分派。该功能直接影响案件进入工作台后的第一小时效率。")
    add_heading(doc, "2. 案件分析代理", 2)
    add_body_paragraph(doc, "在现有 VALUE 能力基础上，升级为多输出形态：争议摘要、双方诉求图谱、证据缺口清单、风险矩阵、可调区间推测、首轮沟通策略和调解推进建议。调解员不再需要从零整理案情。")
    add_heading(doc, "3. 调解策略副驾驶", 2)
    add_body_paragraph(doc, "针对调解前准备、首轮接触、僵局处理、方案比较、协议落地等节点，输出“建议说法、建议动作、注意事项、引用依据、下一步模板”。这类功能更接近 Harvey 式专业工作副驾驶，能显著降低新人上手门槛。")
    add_heading(doc, "4. 协议与文书生成中心", 2)
    add_body_paragraph(doc, "将协议、通知、会议纪要、办案记录、结案摘要等统一纳入模板中心。文书生成不能只做文本填充，而要与案件字段、代理结论、审批流和版本留痕打通，形成真正可复用的交付中心。")
    add_heading(doc, "5. 机构知识工厂", 2)
    add_body_paragraph(doc, "每个结案案件都应进入复盘流程，将有效策略、关键表述、失败原因、协议条款和行业经验写回知识底座。系统由此形成组织记忆，而不只是搜索引擎。")
    add_heading(doc, "6. 负责人管理看板", 2)
    add_body_paragraph(doc, "提供案件漏斗、平均处理时长、超期风险、代理使用率、知识命中率、文书复用率、协议达成率和调解员工作负荷分布，支撑机构负责人做产能管理和流程优化。")

    add_heading(doc, "五、建议的代理体系设计与执行链路", 1)
    add_table(
        doc,
        ["代理名称", "触发时机", "主要输入", "主要输出"],
        [
            ["受理分诊代理", "新案件创建后", "案情描述、材料、来源渠道", "案件标签、优先级、分派建议"],
            ["材料研究代理", "案件进入分析阶段", "上传证据、OCR 文本、历史沟通", "事实摘要、证据缺口、争议焦点"],
            ["策略建议代理", "调解前与沟通中", "双方诉求、风险点、历史回合", "沟通建议、方案路径、话术提示"],
            ["协议生成代理", "达成方向后", "共识条款、支付安排、履约节点", "协议草案、纪要、后续通知"],
            ["复盘沉淀代理", "结案后", "办案记录、结果、调解员备注", "经验卡片、模板优化建议、知识入库条目"],
        ],
        [1.3, 1.2, 1.7, 2.3],
    )
    add_body_paragraph(doc, "这些代理不应被设计成孤立按钮，而应成为案件状态机中的可编排节点。调解员可以接受、编辑、驳回或追问代理结果，系统同时保留引用来源和操作日志。")
    add_table(
        doc,
        ["执行阶段", "运行时动作", "建议存储对象", "失败处理"],
        [
            ["Plan", "生成任务计划与步骤图", "agent_run、agent_step", "允许人工改写计划并重新开始"],
            ["Research", "召回知识、抽取证据、汇总事实", "retrieval_log、evidence_ref", "检索为空时降级为人工补料"],
            ["Work", "调用分析工具、生成策略、整理文书", "tool_call、draft_snapshot", "支持超时重试和步骤重跑"],
            ["Deliver", "输出摘要、文书草案、任务建议", "deliverable、version", "要求人工确认后才能外发"],
            ["Review", "记录采纳率、修改点、复盘结论", "review_note、knowledge_item", "将低质量结果打回模板优化"],
        ],
        [1.0, 1.8, 1.7, 2.0],
    )

    add_heading(doc, "六、对现有项目的具体技术改造建议", 1)
    add_lead_paragraph(doc, "第一步：引入案件工作流引擎", "在 `mediation-workbench/server` 内增加明确的案件状态机、任务表、代理执行记录表和审批流表，把当前分散在 API 和 utils 中的分析动作整合为可追踪的工作流。建议核心表至少包含 `cases`、`case_tasks`、`case_task_runs`、`case_events`、`approvals`。")
    add_lead_paragraph(doc, "第二步：重构 AI 能力入口", "将现有 public AI、case analysis、VALUE skills、chat workflow 统一为“代理执行层”，形成标准输入输出协议。协议需要显式定义 `input_context`、`retrieval_refs`、`tool_results`、`draft_output`、`review_state`，便于后续替换模型或接入多模型路由。")
    add_lead_paragraph(doc, "第三步：建设知识对象模型", "把法规、案例、机构模板、调解记录、行业经验、协议范本、培训资料抽象成统一知识对象。每个对象至少包含 `source_type`、`owner_scope`、`applicable_stage`、`sensitivity_level`、`version`、`embedding_status` 等字段，以支持权限过滤和索引重建。")
    add_lead_paragraph(doc, "第四步：补全组织治理模型", "在用户、角色、机构、案件、材料之外，引入团队、岗位、审批节点、共享空间、审计日志、模板版本和知识贡献度等对象。权限建议采用 RBAC + 资源作用域双层模型，即“角色定义能力，作用域定义可见范围”。")
    add_lead_paragraph(doc, "第五步：建立异步任务与事件总线", "对 OCR、索引、代理执行、文书生成、通知发送等耗时操作，建议统一接入任务队列。早期可用数据库任务表 + worker 进程实现，避免过早引入重型基础设施；当并发上升后，再平滑切换到 Redis / 消息队列。")
    add_lead_paragraph(doc, "第六步：补齐可观测性", "至少建设三类观测数据：业务指标、系统指标、模型指标。业务侧关注案件流转、代理采纳率、协议达成率；系统侧关注接口耗时、队列积压、任务失败率；模型侧关注 token 消耗、检索命中率、引用缺失率、人工回退率。")

    add_heading(doc, "七、建议的核心数据模型", 1)
    add_table(
        doc,
        ["对象", "关键字段", "设计目的", "备注"],
        [
            ["Case", "tenant_id、case_number、stage、owner_id", "描述案件主状态与归属", "作为所有任务和文书的根对象"],
            ["CaseTask", "case_id、task_type、status、assignee_id", "承载人工任务与代理任务", "支持 SLA 和待办聚合"],
            ["AgentRun", "task_id、agent_type、plan_json、status", "记录一次代理执行", "支持重跑、回放与比对"],
            ["EvidenceRef", "run_id、source_id、chunk_id、quote", "记录引用证据链", "保证结果可追溯"],
            ["KnowledgeItem", "scope、tags、stage、version、quality_score", "沉淀机构知识资产", "支持上线后持续优化"],
            ["Deliverable", "case_id、template_id、version、approval_state", "管理协议与文书版本", "连接审批和外发流程"],
        ],
        [1.0, 2.2, 1.7, 1.6],
    )

    add_heading(doc, "八、商业化落地建议", 1)
    add_heading(doc, "1. 产品包装", 2)
    add_body_paragraph(doc, "建议采用“三层产品”包装。基础版解决案件管理与协同，专业版增加代理分析、模板中心和知识库，旗舰版再加入机构治理、数据看板、定制代理和私有化部署能力。")
    add_heading(doc, "2. 收费模型", 2)
    add_body_paragraph(doc, "优先选择“机构年费 + 席位费 + AI 使用包 + 实施服务费”的混合模式。这样既能形成稳定 ARR，也能把前期咨询、知识梳理、模板配置和培训收入纳入商业模型。")
    add_heading(doc, "3. 首批客户选择", 2)
    add_body_paragraph(doc, "优先切入有明确办案流程和管理痛点的调解机构、商会调解中心、行业协会调解中心和园区争议解决中心。这些组织更容易理解内部提效价值，也更愿意采购工作台型产品。")
    add_heading(doc, "4. 交付路径", 2)
    add_body_paragraph(doc, "商业上不建议一开始追求大而全。应先用标准版 SaaS 快速验证“案件周转时间下降、文书生产效率提升、新人培训周期缩短、知识复用提升”四个核心价值，再逐步推出行业模板和机构定制代理。")

    add_heading(doc, "九、分阶段实施路线", 1)
    add_table(
        doc,
        ["阶段", "时间建议", "建设重点", "验收指标"],
        [
            ["阶段一", "0-3 个月", "案件状态机、任务表、代理运行记录、基础检索引用", "单案准备时间下降，代理结果可回放"],
            ["阶段二", "3-6 个月", "文书中心、知识索引流水线、审批流、管理看板", "文书复用率和知识命中率提升"],
            ["阶段三", "6-12 个月", "多租户治理、队列化运行时、定制代理、套餐化能力", "形成可复制销售与续约模型"],
        ],
        [1.0, 1.0, 2.6, 1.9],
    )
    add_body_paragraph(doc, "从实施节奏看，最重要的不是先把模型做得多复杂，而是先把“状态、任务、证据、版本、审批、日志”这六个基础对象打牢。只要这六类对象设计稳定，后续无论是升级模型、替换知识引擎还是拆分服务，系统都能低风险演进。")

    add_heading(doc, "十、最终建议", 1)
    add_lead_paragraph(doc, "战略方向", "把平台从“带 AI 的调解系统”升级为“调解机构工作台 SaaS”，用组织级工作流与知识沉淀形成壁垒。")
    add_lead_paragraph(doc, "产品重点", "优先做案件状态机、专业代理、模板中心、知识工厂和管理看板，而不是继续扩散式增加零散页面。")
    add_lead_paragraph(doc, "商业重点", "围绕机构年费、席位费、AI 使用包和实施服务形成持续收入，先证明内部提效，再做更强的行业复制。")
    add_lead_paragraph(doc, "执行原则", "所有 AI 结论都要可引用、可编辑、可审批、可留痕。调解员负责判断，系统负责准备、执行和沉淀。")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    import sys

    if len(sys.argv) != 2:
        raise SystemExit("Usage: generate_mediation_saas_strategy_doc.py OUTPUT_PATH")
    build_doc(Path(sys.argv[1]).resolve())
